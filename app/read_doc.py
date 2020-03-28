import docx

def ReadingTextDocument(FileName):
    doc= docx.Document(FileName)
    text=[]
    for paragraph in doc.paragraphs:
        text.append (paragraph.text)
    doc= '\n'.join(text)
    return doc

ReadingTextDocument
